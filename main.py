import google.generativeai as genai
import tkinter as tk
import tkcalendar
import tkinter.scrolledtext as st
from tkinter import filedialog

import threading
import json
import re

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime, date, timedelta
from timetable import cut_timetable, datetime_range

#### ADD YOUR API KEY HERE ####
genai.configure(api_key="")

def process_json(json_text):
    result = None
    while True:
        try:
            result = json.loads(json_text)  
            break               
        except Exception as e:
            unexp = int(re.findall(r'\(char (\d+)\)', str(e))[0])
            unesc = json_text.rfind(r'"', 0, unexp)
            json_text = json_text[:unesc] + r'\"' + json_text[unesc+1:]
            closg = json_text.find(r'"', unesc + 2)
            json_text = json_text[:closg] + r'\"' + json_text[closg+1:]

    return result

class QGBot:
    def __init__(self, bot_name):
        self.bot_name = bot_name

        self.generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 64,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }

        self.model = genai.GenerativeModel(
            model_name="gemini-1.5-flash",
            generation_config=self.generation_config,
        )

        
        self.chat_session = self.model.start_chat(history=[])

    def generate_question(self, question_count, topic, school_name, test_name, grade, file_path, status_label):
        t = threading.Thread(target=self.generate_question_t, args=(question_count, topic, school_name, test_name,grade, file_path,status_label, ))
        t.start()

    def generate_question_t(self, question_count, topic, school_name, test_name, grade, file_path,status_label):
        status_label.config(text='ĐANG XUẤT FILE CÂU HỎI....')
        prompt = """Xuất ra """ + question_count + """ câu hỏi trắc nghiệm về chủ đề """ + topic + """ dành cho học sinh lớp """ + grade + """
        Câu hỏi phải được in ra theo định dạng mảng các JSON objects:
        [{
            'question': câu hỏi ở đây,
            'choices': {
                'A': lựa chọn A,
                'B': lựa chọn B,
                'C': lựa chọn C,
                'D': lựa chọn D,
            }
        }]
        """
        text = self.chat_session.send_message(prompt).text
        first_index = text.find('[')
        last_index = text.find(']')
        questions = process_json(text[first_index:last_index + 1])

        document = Document()
        school_p = document.add_heading(f'Trường {school_name}')
        school_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        test_p = document.add_heading(f'Bài kiểm tra {test_name} môn {topic}')
        test_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        document.add_paragraph('');

        document.add_paragraph('Thứ ...... ngày ...... tháng ...... năm .............')
        document.add_paragraph('Họ và tên: ............................................................................................................................')
        document.add_paragraph(f'Lớp: {grade}A...........................................................................')

        q_id = 1
        for question in questions:
            q = f'{q_id}. {question['question']}'
            qp = document.add_paragraph('')
            style = qp.add_run(q)
            style.bold = True
            for choice in question['choices'].keys():
                c = f'\t{choice}. {question['choices'][choice]}'
                cp = document.add_paragraph('')
                style = cp.add_run(c)
                style.bold = True
            q_id += 1
            
        document.save(file_path + f'/{test_name}_{grade}.docx')
        status_label.config(text=(file_path + f'/{test_name}_{grade}.docx saved!'));
         

class QuestionGenApp(tk.Frame):
    def __init__(self, parent):
        tk.Frame.__init__(self, parent)
        self.parent = parent
        self.bot = QGBot('Question Generator')
        self.current_file_path = '';
        self.initUI()
       
    def initUI(self):
        self.parent.title('Question Generator')
        self.default_font = ('consolas', 14)

        self.main_frame = tk.Frame(master=self.parent)

        self.qcount_label = tk.Label(master=self.main_frame, text='Nhập số câu hỏi:', font=self.default_font)
        self.qcount_entry = tk.Entry(master=self.main_frame, font=self.default_font) 
        self.qcount_label.grid(row=0, column=0, padx=15, pady=15)
        self.qcount_entry.grid(row=0, column=1, padx=15, pady=15)
        
        self.topic_label = tk.Label(master=self.main_frame, text='Nhập tên môn học:', font=self.default_font)
        self.topic_entry = tk.Entry(master=self.main_frame, font=self.default_font) 
        self.topic_label.grid(row=1, column=0, padx=15, pady=15)
        self.topic_entry.grid(row=1, column=1, padx=15, pady=15)

        self.school_name_label = tk.Label(master=self.main_frame, text='Nhập tên trường:', font=self.default_font)
        self.school_name_entry = tk.Entry(master=self.main_frame, font=self.default_font) 
        self.school_name_label.grid(row=2, column=0, padx=15, pady=15)
        self.school_name_entry.grid(row=2, column=1, padx=15, pady=15)

        self.testname_label = tk.Label(master=self.main_frame, text='Nhập tên bài kiểm tra:', font=self.default_font)
        self.testname_entry = tk.Entry(master=self.main_frame, font=self.default_font) 
        self.testname_label.grid(row=3, column=0, padx=15, pady=15)
        self.testname_entry.grid(row=3, column=1, padx=15, pady=15)

        self.grade_label = tk.Label(master=self.main_frame, text='Nhập lớp:', font=self.default_font)
        self.grade_entry = tk.Entry(master=self.main_frame, font=self.default_font) 
        self.grade_label.grid(row=4, column=0, padx=15, pady=15)
        self.grade_entry.grid(row=4, column=1, padx=15, pady=15)

        self.file_label = tk.Label(master=self.main_frame, text='Chọn đường dẫn lưu:', font=self.default_font)
        self.file_btn = tk.Button(master=self.main_frame, text='Chọn', font=self.default_font, command=self.open_filediag) 
        self.selectedfile_label = tk.Label(master=self.main_frame, text='', font=self.default_font)
        
        self.file_label.grid(row=5, column=0, padx=15, pady=15)
        self.file_btn.grid(row=5, column=1, padx=15, pady=15)
        self.selectedfile_label.grid(row=6, columnspan=2, padx=15, pady=15);

        self.generate_btn = tk.Button(master=self.main_frame, text='Xuất câu hỏi', font=self.default_font, command=self.on_generate)
        self.generate_btn.grid(row=7, columnspan=2, padx=15, pady=15)

        self.status_label = tk.Label(master=self.main_frame, text='', font=self.default_font)
        self.status_label.grid(row=8, columnspan=2, padx=15, pady=15)
        
        self.exam_day_label = tk.Label(master=self.main_frame, text='Nhập ngày thi:', font=self.default_font)
        self.exam_day_entry = tkcalendar.DateEntry(master=self.main_frame) 
        self.exam_day_label.grid(row=9, column=0, padx=15, pady=15)
        self.exam_day_entry.grid(row=9, column=1, padx=15, pady=15)
        self.generate_timetable_btn = tk.Button(master=self.main_frame, text='Xuất thời khóa biểu', font=self.default_font, command=self.on_generate_timetable)
        self.generate_timetable_btn.grid(row=10, columnspan=2, padx=15, pady=15)

        self.main_frame.pack(padx=10, pady=10)

    def open_filediag(self):
        file_path = filedialog.askdirectory(title="Select dir to save .docx file")
        if file_path:
            self.selectedfile_label.config(text=f"{file_path}")
            self.current_file_path = file_path

    def on_generate_timetable(self):
        if (self.current_file_path != ''):
            timetable = [
                (datetime.today().date(), self.exam_day_entry.get_date(), {"nhiệm vụ:" f"ôn {"bài" if self.topic_entry.get() == "" else self.topic_entry.get()}"}),
            ]
            file = open(self.current_file_path + '/tkb.txt', "w", encoding="utf-8");
            cuts = datetime_range(timetable[0][0], timetable[0][1], timedelta(days=2))
            for sub_timetable in cut_timetable(timetable, cuts):
                for start, end, entry in sub_timetable:
                    s = '%s %s %s\n' % (start.isoformat(), end.isoformat(), entry);
                    print(s)
                    file.write(s)

    def on_generate(self):
        self.status_label.config(text='');
        question_count = self.qcount_entry.get()
        grade = self.grade_entry.get()
        topic = self.topic_entry.get()
        school_name = self.school_name_entry.get()
        test_name = self.testname_entry.get()
        
        self.bot.generate_question(question_count, topic, school_name, test_name, grade, self.current_file_path, self.status_label);
        
        
root = tk.Tk()
app = QuestionGenApp(root)
root.mainloop()